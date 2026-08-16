"""
ingest_knowledge.py

Builds the two files that knowledge_chat.py reads at search time:

    C:\\OllamaAI\\rag\\vector_store.json      <- chunked + embedded text
    C:\\OllamaAI\\knowledge\\metadata.json    <- per-document metadata

Nothing currently creates these two files, which is why search()
in knowledge_chat.py always returns []: the vector store is empty.

USAGE
-----
1. Put your reference files in:  C:\\OllamaAI\\knowledge\\sources\\
   (docx, pdf, and txt are supported)

2. Fill in / edit  C:\\OllamaAI\\knowledge\\docs_manifest.json
   (a starter version is created for you the first time you run this
   script, pre-filled with your current filenames — just edit the
   jurisdiction / document_type / effective_date / status fields)

3. Run:  python ingest_knowledge.py

   Re-run any time you add, remove, or edit a source file — it fully
   rebuilds the vector store from scratch (safe, idempotent).

Requires: pip install python-docx pypdf ollama --break-system-packages
(python-docx and pypdf are only needed if you have .docx / .pdf files)
"""

import json
import os
import sys

import ollama

BASE = r"C:\OllamaAI"

SOURCE_DIR = os.path.join(BASE, "knowledge", "sources")
MANIFEST_FILE = os.path.join(BASE, "knowledge", "docs_manifest.json")
METADATA_FILE = os.path.join(BASE, "knowledge", "metadata.json")
VECTOR_FILE = os.path.join(BASE, "rag", "vector_store.json")

EMBED_MODEL = "nomic-embed-text"

# Chunking parameters (word-based, simple and dependency-free)
CHUNK_WORDS = 350
CHUNK_OVERLAP = 60


# ---------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------

def extract_text_docx(path):
    from docx import Document

    doc = Document(path)
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    # Also pull table content (protocol docs often use tables)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_text_pdf(path):
    from pypdf import PdfReader

    reader = PdfReader(path)
    parts = []

    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text.strip())

    return "\n".join(parts)


def extract_text_txt(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_text(path):
    ext = os.path.splitext(path)[1].lower()

    if ext == ".docx":
        return extract_text_docx(path)
    if ext == ".pdf":
        return extract_text_pdf(path)
    if ext == ".txt":
        return extract_text_txt(path)

    raise ValueError(f"Unsupported file type: {ext}")


# ---------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------

def chunk_text(text, chunk_words=CHUNK_WORDS, overlap=CHUNK_OVERLAP):
    words = text.split()

    if not words:
        return []

    chunks = []
    start = 0

    while start < len(words):
        end = start + chunk_words
        chunk = " ".join(words[start:end])

        if chunk.strip():
            chunks.append(chunk.strip())

        if end >= len(words):
            break

        start = end - overlap

    return chunks


# ---------------------------------------------------------------------
# Manifest / metadata handling
# ---------------------------------------------------------------------

def load_json(path, default):
    if not os.path.exists(path):
        return default

    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def save_json(path, data):
    os.makedirs(os.path.dirname(path), exist_ok=True)

    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)


def ensure_manifest(source_files):
    """
    Create/extend docs_manifest.json with an entry for every source
    file. Existing entries are left untouched so you don't lose edits
    on re-run; only missing files get a placeholder entry added.
    """
    manifest = load_json(MANIFEST_FILE, {})
    changed = False

    for filename in source_files:
        if filename not in manifest:
            manifest[filename] = {
                "title": os.path.splitext(filename)[0].replace("_", " "),
                "jurisdiction": "UNSPECIFIED",
                "document_type": "reference",
                "effective_date": "UNKNOWN",
                "status": "active"
            }
            changed = True

    if changed:
        save_json(MANIFEST_FILE, manifest)
        print(f"Updated manifest: {MANIFEST_FILE}")
        print("  -> Review it and fill in jurisdiction / document_type / "
              "effective_date / status for new entries, then re-run.")

    return manifest


# ---------------------------------------------------------------------
# Main ingestion
# ---------------------------------------------------------------------

def main():
    if not os.path.isdir(SOURCE_DIR):
        os.makedirs(SOURCE_DIR, exist_ok=True)
        print(f"Created empty source folder: {SOURCE_DIR}")
        print("Put your reference docx/pdf/txt files there and re-run.")
        sys.exit(0)

    source_files = sorted(
        f for f in os.listdir(SOURCE_DIR)
        if os.path.splitext(f)[1].lower() in (".docx", ".pdf", ".txt")
    )

    if not source_files:
        print(f"No .docx/.pdf/.txt files found in {SOURCE_DIR}")
        sys.exit(0)

    manifest = ensure_manifest(source_files)

    metadata_out = {"documents": {}}
    vectors_out = []

    for filename in source_files:
        path = os.path.join(SOURCE_DIR, filename)
        info = manifest.get(filename, {})

        print(f"Reading {filename} ...")

        try:
            text = extract_text(path)
        except Exception as e:
            print(f"  SKIPPED ({e})")
            continue

        chunks = chunk_text(text)

        if not chunks:
            print("  No extractable text, skipped.")
            continue

        print(f"  {len(chunks)} chunk(s), embedding...")

        for chunk in chunks:
            response = ollama.embed(
                model=EMBED_MODEL,
                input=chunk
            )

            embedding = response["embeddings"][0]

            vectors_out.append({
                "source": filename,
                "text": chunk,
                "embedding": embedding
            })

        metadata_out["documents"][filename] = {
            "title": info.get("title", filename),
            "jurisdiction": info.get("jurisdiction", "UNSPECIFIED"),
            "document_type": info.get("document_type", "reference"),
            "effective_date": info.get("effective_date", "UNKNOWN"),
            "status": info.get("status", "active")
        }

    save_json(METADATA_FILE, metadata_out)
    save_json(VECTOR_FILE, vectors_out)

    print("")
    print(f"Done. {len(vectors_out)} chunks indexed across "
          f"{len(metadata_out['documents'])} document(s).")
    print(f"Vector store: {VECTOR_FILE}")
    print(f"Metadata:     {METADATA_FILE}")


if __name__ == "__main__":
    main()
